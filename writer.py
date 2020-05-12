from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Cm


class Writer:
    def __init__(self, file_path):
        self._doc = Document(file_path)
        self._tables = self._doc.tables

    def write(self, table_id, array, start_pos):
        _table = self._tables[table_id]
        r_num, c_num = array.shape
        r1st, c1st = start_pos
        row_cursor = 0
        row_prev_tc = None
        for row in _table.rows[r1st:]:
            if row_cursor == r_num:
                break
            col_prev_tc = None
            col_cursor = 0
            unique_col_count = 0
            for cell in row.cells:
                this_tc = cell._tc
                if col_cursor == c_num:
                    break
                if this_tc is col_prev_tc:
                    continue
                col_prev_tc = this_tc
                if unique_col_count < c1st:
                    unique_col_count += 1
                    continue
                if col_cursor == 0:
                    if this_tc is row_prev_tc:
                        break
                    else:
                        row_prev_tc = this_tc
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.text = array[row_cursor, col_cursor]
                col_cursor += 1
            row_cursor += 1
            if this_tc is row_prev_tc:
                continue
        for row in _table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Cm(0.55)

    def _write(self, table_id, array, start_pos):
        _table = self._tables[table_id]
        r_num, c_num = array.shape
        r1st, c1st = start_pos
        rows = _table.rows[r1st:r1st+r_num]
        cols_num = len(rows[0].cells)
        row_cursor = 0
        for row in rows:
            left = 0
            right = 1
            count = 0
            col_cursor = 0
            cells = row.cells
            while right < cols_num and count < c1st + c_num:
                if id(cells[left]) == id(cells[right]):
                    right += 1
                    continue
                else:
                    left = right
                    right += 1
                    count += 1
                    if count < c1st:
                        continue
                    else:
                        cells[left].text = array[row_cursor, col_cursor]
                        col_cursor += 1
            row_cursor += 1

    def _write3(self, table_id, array, start_pos):
        _table = self._tables[table_id]
        r_num, c_num = array.shape
        r1st, c1st = start_pos
        rows = _table.rows[r1st:r1st+r_num]
        row_cursor = 0
        for row in rows:
            col_cursor = 0
            cells = row.cells
            unique_cells = sorted(list(set(cells)), key=cells.index)
            for cell in unique_cells[c1st:c1st+c_num]:
                cell.text = array[row_cursor, col_cursor]
                col_cursor += 1
            row_cursor += 1

    def save(self):
        import time
        time_stamp = time.strftime('%m%d%H%M%S', time.localtime())
        self._doc.save('test/test' + time_stamp + '.docx')


if __name__ == '__main__':
    writer = Writer('data/1抽检原始记录A1级三相外置.docx')
    import numpy as np
    data = np.arange(15*4).reshape(15, 4).astype(str)
    writer.write(3, data, (4, 3))
    writer.save()
