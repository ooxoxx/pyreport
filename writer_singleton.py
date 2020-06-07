from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Cm
import os
import re


class Writer:
    def set_template(self, file_path):
        self._doc = Document(file_path)
        self._file_path = file_path

    def write(self, table_id, array, start_pos):
        if not hasattr(self, '_doc'):
            raise Exception('template unset.')
        _table = self._doc.tables[table_id]
        r_num, c_num = array.shape
        r1st, c1st = start_pos
        row_cursor = 0
        row_prev_tc = None
        for row in _table.rows[r1st:]:
            if row_cursor == r_num:
                break
            col_prev_tc = None
            col_cursor = 0
            unique_col_count = 0
            this_tc = None
            for cell in row.cells:
                this_tc = cell._tc
                if col_cursor == c_num:
                    break
                if this_tc is col_prev_tc:
                    continue
                col_prev_tc = this_tc
                if unique_col_count < c1st:
                    unique_col_count += 1
                    continue
                if col_cursor == 0:
                    if this_tc is row_prev_tc:
                        break
                    else:
                        row_prev_tc = this_tc
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.text = array[row_cursor, col_cursor]
                col_cursor += 1
            if col_cursor == c_num:
                row_cursor += 1
            if this_tc is row_prev_tc:
                continue
        for row in _table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Cm(0.55)

    def save(self):
        import time
        time_stamp = time.strftime('%m%d_%H%M%S_', time.localtime())
        filename = os.path.basename(self._file_path)
        if re.compile(r'^\d').match(filename):
            filename = filename[1:]
        self._doc.save('test/' + time_stamp + filename)


#singleton
writer = Writer()

if __name__ == '__main__':
    writer = Writer()
    writer.set_template('data/1抽检原始记录A1级三相外置.docx')
    import numpy as np
    data = np.arange(15 * 4).reshape(15, 4).astype(str)
    writer.write(3, data, (4, 3))
    writer.save()
